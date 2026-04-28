"""Generate Workflow_Documentation.docx (tables + headings). Run from repo root or anywhere."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(hdr_cells[i], "D9E2F3")
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    document.add_paragraph()


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_path = root / "Workflow_Documentation.docx"

    doc = Document()

    # Title
    t = doc.add_heading("Yelp Discovery App — System Workflow & Roles", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Full-stack workflow: React frontend, FastAPI backend, MongoDB + Kafka. "
        "Includes local search, Yelp Fusion proxy, reviews, favorites, AI assistant, owner dashboard."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    doc.add_heading("1. Actors overview", level=1)
    add_table(
        doc,
        ["Actor", "Account / API", "Typical capabilities"],
        [
            [
                "Guest",
                "No account",
                "Browse Home and Explore; open restaurant details (local + Yelp); read reviews. Log in required for favorites, reviews, chat, profile.",
            ],
            [
                "Diner (user)",
                "POST /auth/signup (role: user)",
                "All guest flows plus Profile, Preferences, Favorites, History; create/update/delete own reviews; AI assistant with saved sessions.",
            ],
            [
                "Owner",
                "POST /auth/owner-signup (role: owner)",
                "Diner-style features where API allows + Owner Dashboard; claim unowned listings; CRUD owned listings; view aggregated reviews on owned restaurants.",
            ],
        ],
    )

    doc.add_heading("2. End-to-end workflows", level=1)
    doc.add_heading("2.1 Discovery & search", level=2)
    doc.add_paragraph(
        "Home (/) → Explore (/explore). Explore combines local DB search and live Yelp search; results drive list/map and detail navigation.",
        style="List Bullet",
    )
    doc.add_paragraph("Local: GET /restaurants — filters (city, cuisine, price, rating, dietary/ambiance, sort, pagination, open_now).", style="List Bullet")
    doc.add_paragraph("Yelp: GET /restaurants/yelp — Fusion proxy for external results.", style="List Bullet")
    doc.add_paragraph()

    doc.add_heading("2.2 Restaurant detail & import", level=2)
    doc.add_paragraph("Local detail: GET /restaurants/{id} (public).", style="List Bullet")
    doc.add_paragraph(
        "Yelp detail: GET /restaurants/yelp/{yelp_id}; optional persist=true (Bearer) upserts to MongoDB.",
        style="List Bullet",
    )
    doc.add_paragraph("Import: POST /restaurants/import-from-yelp with Yelp business id (authenticated).", style="List Bullet")
    doc.add_paragraph("Signed-in local views: POST /users/me/restaurant-views (analytics).", style="List Bullet")
    doc.add_paragraph()

    doc.add_heading("2.3 Reviews", level=2)
    doc.add_paragraph("List: GET /restaurants/{restaurant_id}/reviews.", style="List Bullet")
    doc.add_paragraph("Create: POST /reviews.", style="List Bullet")
    doc.add_paragraph("Update / delete: PUT/DELETE /reviews/{review_id} (author; owner may delete per business rules in service layer).", style="List Bullet")
    doc.add_paragraph()

    doc.add_heading("2.4 Personalization", level=2)
    doc.add_paragraph("Profile: GET/PUT /users/me; POST /users/me/profile-photo.", style="List Bullet")
    doc.add_paragraph("Preferences: GET/PUT /preferences/me.", style="List Bullet")
    doc.add_paragraph("Favorites: GET /favorites/me; POST/DELETE /favorites/{restaurant_id}.", style="List Bullet")
    doc.add_paragraph("History: GET /users/me/history.", style="List Bullet")
    doc.add_paragraph()

    doc.add_heading("2.5 AI dining assistant", level=2)
    doc.add_paragraph("Chat: POST /ai-assistant/chat.", style="List Bullet")
    doc.add_paragraph("Sessions: GET /ai-assistant/sessions; GET /ai-assistant/sessions/{id}; DELETE (requires login).", style="List Bullet")
    doc.add_paragraph()

    doc.add_heading("2.6 Owner-specific API", level=2)
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["GET /owner/dashboard", "Summary metrics, recent activity for owned restaurants"],
            ["GET /owner/restaurants", "List owned listings"],
            ["GET /owner/reviews", "All reviews on owner’s restaurants (filter/sort/pagination)"],
            ["GET /owner/restaurants/{id}/reviews", "Reviews for one owned restaurant"],
            ["POST /restaurants/{id}/claim", "Claim request for unowned listing (authenticated)"],
        ],
    )

    doc.add_heading("2.7 Listing CRUD & uploads", level=2)
    doc.add_paragraph("POST /restaurants — create listing (authenticated).", style="List Bullet")
    doc.add_paragraph("PUT /restaurants/{id} — update if owner of listing.", style="List Bullet")
    doc.add_paragraph("DELETE /restaurants/{id} — delete if owner.", style="List Bullet")
    doc.add_paragraph("POST /uploads/restaurant-photo — image for listing.", style="List Bullet")
    doc.add_paragraph()

    doc.add_heading("3. Frontend routes (reference)", level=1)
    add_table(
        doc,
        ["Path", "Page / behavior"],
        [
            ["/", "HomePage"],
            ["/explore", "ExplorePage — search, map, AI section"],
            ["/login, /signup", "Diner auth"],
            ["/owner/login, /owner/signup", "Owner auth"],
            ["/owner", "OwnerDashboardPage (OwnerRoute)"],
            ["/profile, /preferences, /favorites, /history", "Protected user pages"],
            ["/write-review", "WriteReviewHubPage"],
            ["/restaurants/new", "AddRestaurantPage (protected)"],
            ["/restaurants/:id/edit", "EditRestaurantPage (protected)"],
            ["/restaurants/:id, /restaurants/yelp/:id", "RestaurantDetailsPage"],
            ["/restaurants/:id/review", "WriteReviewPage"],
            ["/restaurants/:id/claim", "ClaimRestaurantPage (protected)"],
        ],
    )

    doc.add_heading("4. Backend route modules (reference)", level=1)
    add_table(
        doc,
        ["Module", "Responsibility"],
        [
            ["auth_routes", "Signup (user/owner), login, GET /auth/me"],
            ["user_routes", "Profile, photo URL, restaurant views, history"],
            ["preference_routes", "Dining preferences GET/PUT"],
            ["restaurant_routes", "Search, Yelp proxy, import, CRUD, claims"],
            ["review_routes", "List by restaurant; CRUD reviews"],
            ["favorite_routes", "List/add/remove favorites"],
            ["owner_routes", "Dashboard, owned restaurants, owner review feeds"],
            ["chat_routes", "AI chat + sessions"],
            ["upload_routes", "Restaurant & profile photo uploads"],
            ["main.py", "App wiring, /health, /docs, static /files"],
        ],
    )

    doc.add_heading("5. Summary", level=1)
    doc.add_paragraph(
        "Guests discover and read public content. Signed-in diners personalize the experience, save favorites, "
        "write reviews, and use the AI assistant with persisted sessions. Owners extend this with a dashboard, "
        "listing management, claims, and visibility into reviews on their businesses—implemented as a React SPA "
        "on FastAPI with MongoDB, optional Yelp Fusion, and Gemini-backed chat."
    )

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
