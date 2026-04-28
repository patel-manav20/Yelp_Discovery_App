"""Convert DEMO_GUIDE.md to DEMO_GUIDE.docx. Run: python scripts/generate_demo_guide_docx.py"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_cell_shading(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_runs_with_bold(paragraph, text: str) -> None:
    """Split on **...** and add bold runs. Escapes nothing."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # strip single *italic* to plain for simplicity
            s = part
            s = re.sub(r"\*([^*]+)\*", r"\1", s)
            paragraph.add_run(s)


def add_paragraph_formatted(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_runs_with_bold(p, text)


def parse_table_lines(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Parse markdown table starting at lines[start]. Returns (rows, next_index)."""
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.match(r"^[\s\-:|]+$", c) for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_md_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs_with_bold(p, val)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "D9E2F3")
    doc.add_paragraph()


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "DEMO_GUIDE.md"
    out_path = root / "DEMO_GUIDE.docx"

    if not md_path.is_file():
        raise SystemExit(f"Missing {md_path}")

    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    i = 0
    code_mode = False
    code_buf: list[str] = []

    def flush_code() -> None:
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        doc.add_paragraph()
        code_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if code_mode:
                flush_code()
                code_mode = False
            else:
                code_mode = True
            i += 1
            continue

        if code_mode:
            code_buf.append(line)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            rows, i = parse_table_lines(lines, i)
            add_md_table(doc, rows)
            continue

        if stripped.startswith("#"):
            level = 0
            rest = stripped
            while rest.startswith("#"):
                level += 1
                rest = rest[1:]
            rest = rest.strip()
            lvl = min(max(level, 1), 3)
            if level == 1:
                h = doc.add_heading(rest, level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(rest, level=lvl - 1)
            i += 1
            continue

        if stripped.startswith("- [ ] "):
            add_paragraph_formatted(doc, stripped[6:].strip(), style="List Bullet")
            i += 1
            continue

        if stripped.startswith("- "):
            add_paragraph_formatted(doc, stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            add_paragraph_formatted(doc, m.group(2).strip(), style="List Number")
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        add_paragraph_formatted(doc, stripped)
        i += 1

    if code_buf:
        flush_code()

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
